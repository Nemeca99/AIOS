#!/usr/bin/env python3
"""
Convert markdown paper to Word document for Research Square submission
"""

import markdown
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
import re


def convert_markdown_to_docx(md_file, docx_file):
    """Convert markdown file to Word document with proper formatting"""

    # Read markdown content
    with open(md_file, "r", encoding="utf-8") as f:
        md_content = f.read()

    # Create Word document
    doc = Document()

    # Set up styles
    styles = doc.styles

    # Title style
    title_style = styles.add_style("CustomTitle", WD_STYLE_TYPE.PARAGRAPH)
    title_style.font.size = Pt(18)
    title_style.font.bold = True
    title_style.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Heading 1 style
    h1_style = styles.add_style("CustomH1", WD_STYLE_TYPE.PARAGRAPH)
    h1_style.font.size = Pt(16)
    h1_style.font.bold = True

    # Heading 2 style
    h2_style = styles.add_style("CustomH2", WD_STYLE_TYPE.PARAGRAPH)
    h2_style.font.size = Pt(14)
    h2_style.font.bold = True

    # Heading 3 style
    h3_style = styles.add_style("CustomH3", WD_STYLE_TYPE.PARAGRAPH)
    h3_style.font.size = Pt(12)
    h3_style.font.bold = True

    # Normal text style
    normal_style = styles.add_style("CustomNormal", WD_STYLE_TYPE.PARAGRAPH)
    normal_style.font.size = Pt(11)

    # Split content into lines
    lines = md_content.split("\n")

    for line in lines:
        line = line.strip()

        if not line:
            # Empty line - add paragraph break
            doc.add_paragraph()
            continue

        # Handle headers
        if line.startswith("# "):
            # Main title
            title = line[2:]
            p = doc.add_paragraph(title, style="CustomTitle")
            continue

        elif line.startswith("## "):
            # Heading 1
            heading = line[3:]
            p = doc.add_paragraph(heading, style="CustomH1")
            continue

        elif line.startswith("### "):
            # Heading 2
            heading = line[4:]
            p = doc.add_paragraph(heading, style="CustomH2")
            continue

        elif line.startswith("#### "):
            # Heading 3
            heading = line[5:]
            p = doc.add_paragraph(heading, style="CustomH3")
            continue

        # Handle horizontal rules
        elif line.startswith("---"):
            # Add a horizontal line (represented as a paragraph with underscores)
            p = doc.add_paragraph("_" * 50, style="CustomNormal")
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            continue

        # Handle bold text
        elif "**" in line:
            # Convert markdown bold to Word bold
            line = re.sub(r"\*\*(.*?)\*\*", r"\1", line)
            p = doc.add_paragraph(line, style="CustomNormal")
            # Make the text bold
            for run in p.runs:
                run.bold = True
            continue

        # Handle lists
        elif line.startswith("- "):
            # Bullet point
            content = line[2:]
            p = doc.add_paragraph(content, style="CustomNormal")
            p.style = "List Bullet"
            continue

        elif line.startswith("1. "):
            # Numbered list
            content = line[3:]
            p = doc.add_paragraph(content, style="CustomNormal")
            p.style = "List Number"
            continue

        # Regular paragraph
        else:
            p = doc.add_paragraph(line, style="CustomNormal")

    # Save the document
    doc.save(docx_file)
    print(f"Successfully converted {md_file} to {docx_file}")


if __name__ == "__main__":
    # Convert the paper
    convert_markdown_to_docx(
        "paper/TEMPORAL_MECHANICS_PAPER.md", "TEMPORAL_MECHANICS_PAPER.docx"
    )
